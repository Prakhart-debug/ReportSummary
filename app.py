import os
import pandas as pd
import docx
import numpy as np
import matplotlib.pyplot as plt
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics import silhouette_score
from sklearn.cluster import KMeans
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import nltk
import streamlit as st
from io import BytesIO
from groq import Groq

# Ensure nltk resources are downloaded
nltk.download('stopwords')
nltk.download('punkt')

# Groq API Key (ensure this is set properly)
GROQ_API_KEY = "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"
client = Groq(api_key=GROQ_API_KEY)

# Model options
MODELS = {
    "LLaMA3 8b": "llama3-8b-8192",
    "LLaMA3 70b": "llama3-70b-8192",
    "LLaMa3 Groq 8b Tool Use": "llama3-groq-8b-8192-tool-use-preview",
    "LLaMA3 Groq 70b Tool Use": "llama3-groq-70b-8192-tool-use-preview",
    "Mixtral 8x7b": "mixtral-8x7b-32768",
    "Gemma 7b": "gemma-7b-it",
    "Gemma2 9b": "gemma2-9b-it"
}

# Define helper functions
def read_docx(file_path):
    doc = docx.Document(file_path)
    content = []
    for para in doc.paragraphs:
        content.append(para.text)
    return ' '.join(content)

def write_docx(file_path, content):
    doc = docx.Document()
    doc.add_paragraph(content)
    doc.save(file_path)

def optimal_number_of_clusters(tfidf_matrix, max_clusters=10):
    n_samples = tfidf_matrix.shape[0]
    max_clusters = min(max_clusters, n_samples - 1)
    
    if n_samples < 3:
        return 1
    
    if n_samples < max_clusters:
        max_clusters = n_samples - 1
    
    K = range(2, max_clusters + 1)
    
    if n_samples >= 4:
        silhouette_scores = []
        for k in K:
            kmeans = KMeans(n_clusters=k, random_state=42, n_init=10)
            kmeans.fit(tfidf_matrix)
            score = silhouette_score(tfidf_matrix, kmeans.labels_)
            silhouette_scores.append(score)
        
        return silhouette_scores.index(max(silhouette_scores)) + 2
    else:
        inertias = []
        for k in K:
            kmeans = KMeans(n_clusters=k, random_state=42, n_init=10)
            kmeans.fit(tfidf_matrix)
            inertias.append(kmeans.inertia_)
        
        diffs = np.diff(inertias)
        elbow_point = np.argmin(diffs) + 2
        return elbow_point

def cluster_documents(documents):
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform(documents)
    
    n_clusters = optimal_number_of_clusters(tfidf_matrix)
    
    kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
    kmeans.fit(tfidf_matrix)
    
    return kmeans.labels_, vectorizer

def consolidate_content(content, model_id):
    response = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": f"Please consolidate and summarize the following content, removing any duplicate information while keeping all unique and relevant information: Output should contain the heading, followed a professional summary, make sure that there are multiple headlines, and the words should be around 800-1000\n\n{content}"
            }
        ],
        model=model_id
    )
    return response.choices[0].message.content.strip()

def get_cluster_name(cluster_docs, vectorizer):
    combined_text = ' '.join(cluster_docs)
    
    stop_words = set(stopwords.words('english'))
    word_tokens = word_tokenize(combined_text.lower())
    filtered_words = [w for w in word_tokens if w.isalnum() and w not in stop_words]
    
    word_freq = nltk.FreqDist(filtered_words)
    
    for word, _ in word_freq.most_common():
        if word in vectorizer.get_feature_names_out():
            return word.capitalize()
    
    return "Cluster"

def create_report_file(content, cluster_name):
    doc = docx.Document()
    doc.add_paragraph(content)
    file_path = f"{cluster_name}_Consolidated_Output.docx"
    doc.save(file_path)
    return file_path

# Streamlit UI
st.title("Document Clustering and Consolidation")

uploaded_files = st.file_uploader("Choose .docx files", type="docx", accept_multiple_files=True)
selected_model = st.selectbox("Select Model", list(MODELS.keys()))

if st.button("Process Documents"):
    if uploaded_files:
        documents = [read_docx(file) for file in uploaded_files]
        file_names = [file.name for file in uploaded_files]
        
        labels, vectorizer = cluster_documents(documents)
        
        reports = []
        for cluster_id in set(labels):
            cluster_docs = [documents[i] for i in range(len(documents)) if labels[i] == cluster_id]
            cluster_files = [file_names[i] for i in range(len(file_names)) if labels[i] == cluster_id]
            
            combined_content = ' '.join(cluster_docs)
            consolidated_content = consolidate_content(combined_content, MODELS[selected_model])
            cluster_name = get_cluster_name(cluster_docs, vectorizer)
            
            report_file = create_report_file(consolidated_content, cluster_name)
            reports.append(report_file)
            
            st.write(f"Cluster {cluster_id} consolidated into: {report_file}")
            st.write(f"Files in this cluster: {', '.join(cluster_files)}")

            # Provide download link
            with open(report_file, "rb") as f:
                st.download_button(
                    label="Download Report",
                    data=f,
                    file_name=report_file,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.error("Please upload some .docx files.")

